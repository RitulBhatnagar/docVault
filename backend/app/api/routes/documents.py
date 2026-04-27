import hashlib
import io
import uuid
from datetime import date
from pathlib import Path
from typing import Any, Literal

from fastapi import APIRouter, BackgroundTasks, Form, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy import literal_column, or_
from sqlalchemy.orm import selectinload
from sqlmodel import col, func, select

from app.api.deps import CurrentUser, SessionDep
from app.core.config import settings
from app.services import storage as storage_svc
from app.models import (
    BulkDeleteRequest,
    Document,
    DocumentGroup,
    DocumentGroupsPublic,
    DocumentPublic,
    DocumentTag,
    DocumentVersion,
    DocumentVersionPublic,
    DocumentWithVersions,
    DocumentsPublic,
    Message,
    StorageStats,
    Tag,
    TagPublic,
)

router = APIRouter(prefix="/documents", tags=["documents"])

UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


_MAX_TEXT_CHARS = 100_000
_MAX_PDF_PAGES = 50
_MAX_XLSX_ROWS = 200
_MAX_EXTRACT_BYTES = 2 * 1024 * 1024  # skip extraction for files > 2 MB


def _extract_text(content: bytes, filename: str) -> str | None:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if len(content) > _MAX_EXTRACT_BYTES and ext in ("xlsx", "xls"):
        return None
    text: str | None = None
    try:
        if ext == "pdf":
            from pypdf import PdfReader  # noqa: PLC0415
            reader = PdfReader(io.BytesIO(content))
            pages = reader.pages[:_MAX_PDF_PAGES]
            text = " ".join(page.extract_text() or "" for page in pages)
        elif ext in ("docx", "doc"):
            from docx import Document as DocxDocument  # noqa: PLC0415
            doc = DocxDocument(io.BytesIO(content))
            text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in ("txt", "md", "csv"):
            text = content[:_MAX_TEXT_CHARS].decode("utf-8", errors="ignore")
        elif ext in ("xlsx", "xls"):
            import openpyxl  # noqa: PLC0415
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            parts: list[str] = []
            for ws in wb.worksheets:
                for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    if row_idx >= _MAX_XLSX_ROWS:
                        break
                    parts.extend(str(v) for v in row if v is not None)
            text = " ".join(parts)
    except Exception:
        return None
    if not text:
        return None
    text = text.replace("\x00", "")
    return text[:_MAX_TEXT_CHARS]


def _check_duplicate(session: Any, owner_id: uuid.UUID, sha256: str) -> None:
    dup = session.exec(
        select(DocumentVersion)
        .join(Document, DocumentVersion.document_id == Document.id)  # type: ignore[arg-type]
        .where(Document.owner_id == owner_id)
        .where(DocumentVersion.sha256 == sha256)
    ).first()
    if dup:
        dup_doc = session.get(Document, dup.document_id)
        title = dup_doc.title if dup_doc else "unknown"
        raise HTTPException(
            status_code=409,
            detail=f'Duplicate: file already exists as "{title}"',
        )


@router.post("/", response_model=DocumentPublic)
async def upload_document(
    *,
    background_tasks: BackgroundTasks,
    session: SessionDep,
    current_user: CurrentUser,
    title: str = Form(...),
    creator: str = Form(...),
    format: str = Form(...),
    subject: str | None = Form(default=None),
    file: UploadFile,
) -> Any:
    from app.services.ocr import needs_ocr, run_ocr_for_version  # noqa: PLC0415

    content = await file.read()
    sha256 = hashlib.sha256(content).hexdigest()

    _check_duplicate(session, current_user.id, sha256)

    doc = Document(
        title=title,
        creator=creator,
        format=format,
        subject=subject,
        owner_id=current_user.id,
    )
    session.add(doc)
    session.flush()

    filename = file.filename or "unknown"
    object_key = f"{doc.id}_v1_{filename}"
    if settings.r2_enabled:
        storage_svc.upload_bytes(object_key, content, _mime_for_filename(filename))
        stored_path = storage_svc.make_r2_path(object_key)
    else:
        file_path = UPLOAD_DIR / object_key
        file_path.write_bytes(content)
        stored_path = str(file_path)

    content_text = _extract_text(content, filename)
    ocr_needed = needs_ocr(content_text, filename)

    version = DocumentVersion(
        document_id=doc.id,
        version_number=1,
        sha256=sha256,
        file_path=stored_path,
        original_filename=filename,
        file_size=len(content),
        content_text=content_text,
        ocr_status="pending" if ocr_needed else None,
    )
    session.add(version)
    session.commit()
    session.refresh(doc)

    if ocr_needed:
        background_tasks.add_task(
            run_ocr_for_version,
            version.id,
            str(settings.SQLALCHEMY_DATABASE_URI),
        )

    return doc


@router.get("/stats", response_model=StorageStats)
def get_storage_stats(
    *,
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    owner_filter = (
        [] if current_user.is_superuser else [Document.owner_id == current_user.id]
    )
    doc_count = session.exec(
        select(func.count()).select_from(Document).where(*owner_filter)
    ).one()
    ver_stats = session.exec(
        select(func.count(), func.coalesce(func.sum(DocumentVersion.file_size), 0))
        .select_from(DocumentVersion)
        .join(Document, DocumentVersion.document_id == Document.id)  # type: ignore[arg-type]
        .where(*owner_filter)
    ).one()
    return StorageStats(
        document_count=doc_count,
        version_count=ver_stats[0],
        total_size_bytes=ver_stats[1],
    )


@router.get("/search", response_model=DocumentsPublic)
def search_documents(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    q: str,
    skip: int = 0,
    limit: int = 100,
) -> Any:
    tsquery = func.plainto_tsquery("english", q)

    meta_match = literal_column("document.metadata_tsv").op("@@")(tsquery)

    content_subq = (
        select(DocumentVersion.document_id)
        .where(literal_column("documentversion.content_tsv").op("@@")(tsquery))
        .subquery()
    )

    where_clause = or_(meta_match, Document.id.in_(content_subq))  # type: ignore[attr-defined]

    count_stmt = (
        select(func.count())
        .select_from(Document)
        .where(Document.owner_id == current_user.id)
        .where(where_clause)
    )
    count = session.exec(count_stmt).one()

    stmt = (
        select(Document)
        .where(Document.owner_id == current_user.id)
        .where(where_clause)
        .offset(skip)
        .limit(limit)
    )
    docs = session.exec(stmt).all()
    return DocumentsPublic(data=list(docs), count=count)


@router.get("/", response_model=DocumentsPublic)
def list_documents(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = 0,
    limit: int = 100,
    sort_by: Literal["created_at", "title", "format", "creator"] = Query(default="created_at"),
    sort_order: Literal["asc", "desc"] = Query(default="desc"),
    format: str | None = Query(default=None),
    date_from: date | None = Query(default=None),
    date_to: date | None = Query(default=None),
    tag_id: uuid.UUID | None = Query(default=None),
) -> Any:
    sort_col_map = {
        "created_at": Document.created_at,
        "title": Document.title,
        "format": Document.format,
        "creator": Document.creator,
    }
    sort_col = sort_col_map[sort_by]
    order_expr = col(sort_col).asc() if sort_order == "asc" else col(sort_col).desc()

    count_stmt = select(func.count()).select_from(Document)
    docs_stmt = select(Document).options(selectinload(Document.tags))  # type: ignore[arg-type]

    if not current_user.is_superuser:
        count_stmt = count_stmt.where(Document.owner_id == current_user.id)
        docs_stmt = docs_stmt.where(Document.owner_id == current_user.id)
    if format:
        count_stmt = count_stmt.where(col(Document.format).ilike(f"%{format}%"))
        docs_stmt = docs_stmt.where(col(Document.format).ilike(f"%{format}%"))
    if date_from:
        count_stmt = count_stmt.where(func.date(Document.created_at) >= date_from)
        docs_stmt = docs_stmt.where(func.date(Document.created_at) >= date_from)
    if date_to:
        count_stmt = count_stmt.where(func.date(Document.created_at) <= date_to)
        docs_stmt = docs_stmt.where(func.date(Document.created_at) <= date_to)
    if tag_id:
        count_stmt = count_stmt.where(
            Document.id.in_(select(DocumentTag.document_id).where(DocumentTag.tag_id == tag_id))  # type: ignore[attr-defined]
        )
        docs_stmt = docs_stmt.where(
            Document.id.in_(select(DocumentTag.document_id).where(DocumentTag.tag_id == tag_id))  # type: ignore[attr-defined]
        )

    count = session.exec(count_stmt).one()
    docs = session.exec(docs_stmt.order_by(order_expr).offset(skip).limit(limit)).all()
    return DocumentsPublic(data=list(docs), count=count)


@router.get("/check-title", response_model=DocumentPublic | None)
def check_document_title(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    title: str = Query(..., min_length=1),
) -> Any:
    doc = session.exec(
        select(Document)
        .options(selectinload(Document.tags))  # type: ignore[arg-type]
        .where(
            Document.owner_id == current_user.id,
            func.lower(Document.title) == title.strip().lower(),
        )
    ).first()
    return doc


@router.get("/groups", response_model=DocumentGroupsPublic)
def get_document_groups(
    *,
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    from app.services.grouping import derive_cluster_key  # noqa: PLC0415

    docs = session.exec(
        select(Document)
        .options(selectinload(Document.tags))  # type: ignore[arg-type]
        .where(Document.owner_id == current_user.id)
        .order_by(Document.created_at.desc())
    ).all()

    doc_ids = [d.id for d in docs]

    if doc_ids:
        max_ver_subq = (
            select(
                DocumentVersion.document_id,
                func.max(DocumentVersion.version_number).label("max_ver"),
            )
            .where(col(DocumentVersion.document_id).in_(doc_ids))
            .group_by(DocumentVersion.document_id)
            .subquery()
        )
        latest_versions = session.exec(
            select(DocumentVersion).join(
                max_ver_subq,
                (DocumentVersion.document_id == max_ver_subq.c.document_id)
                & (DocumentVersion.version_number == max_ver_subq.c.max_ver),
            )
        ).all()
        latest_ver_map: dict = {v.document_id: v for v in latest_versions}
    else:
        latest_ver_map = {}

    groups_map: dict[str, dict] = {}
    for doc in docs:
        latest_ver = latest_ver_map.get(doc.id)
        key, label, kind = derive_cluster_key(doc, latest_ver)
        if key not in groups_map:
            groups_map[key] = {"key": key, "label": label, "kind": kind, "docs": []}
        groups_map[key]["docs"].append(doc)

    groups: list[DocumentGroup] = []
    other_group: DocumentGroup | None = None

    for g in groups_map.values():
        group = DocumentGroup(
            key=g["key"],
            label=g["label"],
            kind=g["kind"],
            count=len(g["docs"]),
            docs=g["docs"],
        )
        if g["key"] == "other":
            other_group = group
        else:
            groups.append(group)

    groups.sort(key=lambda x: x.count, reverse=True)
    if other_group:
        groups.append(other_group)

    return DocumentGroupsPublic(groups=groups, total=len(docs))


@router.get("/{id}", response_model=DocumentWithVersions)
def get_document(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    versions = session.exec(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == id)
        .order_by(DocumentVersion.version_number)
    ).all()
    return DocumentWithVersions(
        **doc.model_dump(),
        versions=[DocumentVersionPublic.model_validate(v) for v in versions],
    )


@router.post("/{id}/versions", response_model=DocumentVersionPublic)
async def upload_new_version(
    *,
    background_tasks: BackgroundTasks,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    file: UploadFile,
) -> Any:
    from app.services.ocr import needs_ocr, run_ocr_for_version  # noqa: PLC0415

    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    new_ext = file.filename.rsplit(".", 1)[-1].lower() if file.filename and "." in file.filename else ""
    if new_ext and new_ext.lower() != doc.format.lower():
        raise HTTPException(
            status_code=422,
            detail=f"Format mismatch: document is '{doc.format}', uploaded file is '{new_ext}'",
        )

    content = await file.read()
    sha256 = hashlib.sha256(content).hexdigest()

    _check_duplicate(session, current_user.id, sha256)

    max_version = session.exec(
        select(func.max(DocumentVersion.version_number)).where(
            DocumentVersion.document_id == id
        )
    ).one()
    next_version = (max_version or 0) + 1

    filename = file.filename or "unknown"
    object_key = f"{id}_v{next_version}_{filename}"
    if settings.r2_enabled:
        storage_svc.upload_bytes(object_key, content, _mime_for_filename(filename))
        stored_path = storage_svc.make_r2_path(object_key)
    else:
        file_path = UPLOAD_DIR / object_key
        file_path.write_bytes(content)
        stored_path = str(file_path)

    content_text = _extract_text(content, filename)
    ocr_needed = needs_ocr(content_text, filename)

    version = DocumentVersion(
        document_id=id,
        version_number=next_version,
        sha256=sha256,
        file_path=stored_path,
        original_filename=filename,
        file_size=len(content),
        content_text=content_text,
        ocr_status="pending" if ocr_needed else None,
    )
    session.add(version)
    session.commit()
    session.refresh(version)

    if ocr_needed:
        background_tasks.add_task(
            run_ocr_for_version,
            version.id,
            str(settings.SQLALCHEMY_DATABASE_URI),
        )

    return version


@router.get("/{id}/versions", response_model=list[DocumentVersionPublic])
def list_versions(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    return session.exec(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == id)
        .order_by(DocumentVersion.version_number)
    ).all()


MIME_MAP: dict[str, str] = {
    "pdf":  "application/pdf",
    "png":  "image/png",
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "gif":  "image/gif",
    "webp": "image/webp",
    "svg":  "image/svg+xml",
    "txt":  "text/plain",
    "md":   "text/plain",
    "csv":  "text/plain",
    # video
    "mp4":  "video/mp4",
    "webm": "video/webm",
    "ogv":  "video/ogg",
    "mov":  "video/quicktime",
    "avi":  "video/x-msvideo",
    "mkv":  "video/x-matroska",
    # audio
    "mp3":  "audio/mpeg",
    "wav":  "audio/wav",
    "ogg":  "audio/ogg",
    "flac": "audio/flac",
    "m4a":  "audio/mp4",
    "aac":  "audio/aac",
}

_HTML_STYLE = (
    "<style>body{font-family:sans-serif;font-size:13px;padding:12px;margin:0}"
    "table{border-collapse:collapse;width:100%}"
    "th{background:#f3f4f6;font-weight:600;text-align:left}"
    "th,td{border:1px solid #d1d5db;padding:6px 10px;white-space:nowrap}"
    "tr:nth-child(even){background:#f9fafb}"
    "h3{margin:16px 0 6px;font-size:14px;color:#374151}"
    "</style>"
)


def _xlsx_to_html(path: Any) -> str:
    import openpyxl  # noqa: PLC0415
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = [f"<html><head>{_HTML_STYLE}</head><body>"]
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"<h3>{sheet_name}</h3><table>")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            tag = "th" if i == 0 else "td"
            cells = "".join(f"<{tag}>{'' if v is None else v}</{tag}>" for v in row)
            parts.append(f"<tr>{cells}</tr>")
        parts.append("</table>")
    parts.append("</body></html>")
    return "".join(parts)


def _docx_to_html(path: Any) -> str:
    from docx import Document as DocxDocument  # noqa: PLC0415
    doc = DocxDocument(path)
    parts = [f"<html><head>{_HTML_STYLE}</head><body>"]
    for para in doc.paragraphs:
        if not para.text.strip():
            parts.append("<br>")
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            level = style.replace("Heading ", "").strip() or "2"
            parts.append(f"<h{level}>{para.text}</h{level}>")
        else:
            parts.append(f"<p>{para.text}</p>")
    parts.append("</body></html>")
    return "".join(parts)


def _mime_for_filename(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return MIME_MAP.get(ext, "application/octet-stream")


def _preview_response(stored_path: str, filename: str) -> Any:
    from fastapi.responses import StreamingResponse  # noqa: PLC0415
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if storage_svc.is_r2_key(stored_path):
        data = storage_svc.get_file_bytes(stored_path)
        if ext in ("xlsx", "xls"):
            return HTMLResponse(_xlsx_to_html(io.BytesIO(data)))  # type: ignore[arg-type]
        if ext in ("docx", "doc"):
            return HTMLResponse(_docx_to_html(io.BytesIO(data)))  # type: ignore[arg-type]
        return StreamingResponse(io.BytesIO(data), media_type=_mime_for_filename(filename))
    file_path = Path(stored_path)
    if ext in ("xlsx", "xls"):
        return HTMLResponse(_xlsx_to_html(file_path))
    if ext in ("docx", "doc"):
        return HTMLResponse(_docx_to_html(file_path))
    return FileResponse(path=str(file_path), media_type=_mime_for_filename(filename))


@router.get("/{id}/preview")
def preview_latest(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    version = session.exec(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == id)
        .order_by(DocumentVersion.version_number.desc())  # type: ignore[attr-defined]
        .limit(1)
    ).first()
    if not version:
        raise HTTPException(status_code=404, detail="No versions found")

    if storage_svc.is_r2_key(version.file_path):
        return _preview_response(version.file_path, version.original_filename)
    file_path = Path(version.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return _preview_response(str(file_path), version.original_filename)


@router.get("/{id}/versions/{version_id}/preview")
def preview_version(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    version_id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    version = session.get(DocumentVersion, version_id)
    if not version or version.document_id != id:
        raise HTTPException(status_code=404, detail="Version not found")

    if storage_svc.is_r2_key(version.file_path):
        return _preview_response(version.file_path, version.original_filename)
    file_path = Path(version.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return _preview_response(str(file_path), version.original_filename)


@router.get("/{id}/download")
def download_latest(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    version = session.exec(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == id)
        .order_by(DocumentVersion.version_number.desc())  # type: ignore[attr-defined]
        .limit(1)
    ).first()
    if not version:
        raise HTTPException(status_code=404, detail="No versions found")

    if storage_svc.is_r2_key(version.file_path):
        from fastapi.responses import StreamingResponse  # noqa: PLC0415
        data = storage_svc.get_file_bytes(version.file_path)
        headers = {"Content-Disposition": f'attachment; filename="{version.original_filename}"'}
        return StreamingResponse(io.BytesIO(data), media_type="application/octet-stream", headers=headers)
    file_path = Path(version.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(
        path=str(file_path),
        filename=version.original_filename,
        media_type="application/octet-stream",
    )


@router.get("/{id}/versions/{version_id}/download")
def download_version(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    version_id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    version = session.get(DocumentVersion, version_id)
    if not version or version.document_id != id:
        raise HTTPException(status_code=404, detail="Version not found")

    if storage_svc.is_r2_key(version.file_path):
        from fastapi.responses import StreamingResponse  # noqa: PLC0415
        data = storage_svc.get_file_bytes(version.file_path)
        headers = {"Content-Disposition": f'attachment; filename="{version.original_filename}"'}
        return StreamingResponse(io.BytesIO(data), media_type="application/octet-stream", headers=headers)
    file_path = Path(version.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(
        path=str(file_path),
        filename=version.original_filename,
        media_type="application/octet-stream",
    )


@router.delete("/bulk", response_model=Message)
def bulk_delete_documents(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    body: BulkDeleteRequest,
) -> Any:
    deleted = 0
    for doc_id in body.ids:
        doc = session.get(Document, doc_id)
        if not doc:
            continue
        if not current_user.is_superuser and doc.owner_id != current_user.id:
            continue
        versions = session.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)).all()
        for v in versions:
            storage_svc.delete_file(v.file_path)
        session.delete(doc)
        deleted += 1
    session.commit()
    return Message(message=f"Deleted {deleted} document(s)")


@router.delete("/{id}", response_model=Message)
def delete_document(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    versions = session.exec(select(DocumentVersion).where(DocumentVersion.document_id == id)).all()
    for v in versions:
        storage_svc.delete_file(v.file_path)
    session.delete(doc)
    session.commit()
    return Message(message="Document deleted successfully")


# ---------- Tags ----------

@router.get("/{id}/tags", response_model=list[TagPublic])
def list_document_tags(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    return doc.tags


@router.post("/{id}/tags", response_model=TagPublic)
def add_tag_to_document(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    name: str = Query(min_length=1, max_length=50),
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    tag = session.exec(
        select(Tag).where(Tag.owner_id == current_user.id).where(col(Tag.name) == name)
    ).first()
    if not tag:
        tag = Tag(name=name, owner_id=current_user.id)
        session.add(tag)
        session.flush()

    already_linked = session.get(DocumentTag, (id, tag.id))
    if not already_linked:
        session.add(DocumentTag(document_id=id, tag_id=tag.id))
    session.commit()
    session.refresh(tag)
    return tag


@router.delete("/{id}/tags/{tag_id}", response_model=Message)
def remove_tag_from_document(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    id: uuid.UUID,
    tag_id: uuid.UUID,
) -> Any:
    doc = session.get(Document, id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not current_user.is_superuser and doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    link = session.get(DocumentTag, (id, tag_id))
    if not link:
        raise HTTPException(status_code=404, detail="Tag not linked to document")
    session.delete(link)
    session.commit()
    return Message(message="Tag removed")


@router.get("/tags/all", response_model=list[TagPublic])
def list_user_tags(
    *,
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    return session.exec(
        select(Tag).where(Tag.owner_id == current_user.id).order_by(col(Tag.name))
    ).all()


# ---------- OCR Backfill (superuser only) ----------

class OcrBackfillQueued(Message):
    queued: int


class OcrBackfillStatus(Message):
    total_pdf_versions: int
    pending: int
    processing: int
    done: int
    failed: int
    not_applicable: int


@router.post("/ocr/backfill", response_model=OcrBackfillQueued)
def ocr_backfill(
    *,
    background_tasks: BackgroundTasks,
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    from app.services.ocr import run_ocr_backfill  # noqa: PLC0415

    if not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Superuser only")

    candidates = session.exec(
        select(DocumentVersion).where(
            DocumentVersion.ocr_status.is_(None),  # type: ignore[union-attr]
            col(DocumentVersion.original_filename).ilike("%.pdf"),
            DocumentVersion.content_text.is_(None),  # type: ignore[union-attr]
        )
    ).all()

    for v in candidates:
        v.ocr_status = "pending"
        session.add(v)
    session.commit()

    background_tasks.add_task(run_ocr_backfill, str(settings.SQLALCHEMY_DATABASE_URI))

    return OcrBackfillQueued(message="OCR backfill started", queued=len(candidates))


@router.get("/ocr/backfill/status", response_model=OcrBackfillStatus)
def ocr_backfill_status(
    *,
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    if not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Superuser only")

    pdf_versions = session.exec(
        select(DocumentVersion).where(
            col(DocumentVersion.original_filename).ilike("%.pdf")
        )
    ).all()

    counts: dict[str, int] = {"pending": 0, "processing": 0, "done": 0, "failed": 0, "none": 0}
    for v in pdf_versions:
        key = v.ocr_status if v.ocr_status in counts else "none"
        counts[key] += 1

    return OcrBackfillStatus(
        message="ok",
        total_pdf_versions=len(pdf_versions),
        pending=counts["pending"],
        processing=counts["processing"],
        done=counts["done"],
        failed=counts["failed"],
        not_applicable=counts["none"],
    )