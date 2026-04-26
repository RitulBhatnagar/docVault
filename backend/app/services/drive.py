"""Google Drive import service — token management, Drive API helpers, import orchestration."""

from __future__ import annotations

import hashlib
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any

from sqlmodel import Session, select

if TYPE_CHECKING:
    from cryptography.fernet import Fernet  # noqa: F401
    from google.oauth2.credentials import Credentials  # noqa: F401

from app.core.config import settings
from app.core.db import engine
from app.models import (
    Document,
    DocumentTag,
    DocumentVersion,
    DriveConnection,
    DriveImportJob,
    Tag,
)

UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

_GOOGLE_APPS_PREFIX = "application/vnd.google-apps."
_DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]


# ---------- Token encryption ----------


def _fernet() -> Fernet:
    from cryptography.fernet import Fernet  # noqa: PLC0415

    key = settings.DRIVE_TOKEN_ENCRYPTION_KEY
    if not key:
        raise RuntimeError("DRIVE_TOKEN_ENCRYPTION_KEY not configured")
    return Fernet(key.encode() if isinstance(key, str) else key)


def encrypt_token(token: str) -> str:
    return _fernet().encrypt(token.encode()).decode()


def decrypt_token(token: str) -> str:
    return _fernet().decrypt(token.encode()).decode()


# ---------- Google credentials ----------


def build_credentials(connection: DriveConnection) -> Credentials:
    from google.auth.transport.requests import Request  # noqa: PLC0415
    from google.oauth2.credentials import Credentials  # noqa: PLC0415

    creds = Credentials(
        token=decrypt_token(connection.access_token),
        refresh_token=decrypt_token(connection.refresh_token),
        token_uri="https://oauth2.googleapis.com/token",
        client_id=settings.GOOGLE_CLIENT_ID,
        client_secret=settings.GOOGLE_CLIENT_SECRET,
        scopes=_DRIVE_SCOPES,
        expiry=connection.token_expiry.replace(tzinfo=None),
    )
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with Session(engine) as session:
            conn = session.get(DriveConnection, connection.id)
            if conn:
                conn.access_token = encrypt_token(creds.token)
                conn.token_expiry = (
                    creds.expiry.replace(tzinfo=timezone.utc)
                    if creds.expiry
                    else connection.token_expiry
                )
                session.add(conn)
                session.commit()
    return creds


# ---------- Drive API helpers ----------


def list_folders(creds: Credentials) -> list[dict[str, Any]]:
    from googleapiclient.discovery import build  # noqa: PLC0415

    service = build("drive", "v3", credentials=creds)
    results = (
        service.files()
        .list(
            q="mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id,name)",
            pageSize=100,
        )
        .execute()
    )
    sub_folders = results.get("files", [])
    return [{"id": "root", "name": "My Drive (root)"}] + sub_folders


def list_importable_files(creds: Credentials, folder_id: str) -> list[dict[str, Any]]:
    from googleapiclient.discovery import build  # noqa: PLC0415

    service = build("drive", "v3", credentials=creds)
    results = (
        service.files()
        .list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id,name,mimeType,size)",
            pageSize=1000,
        )
        .execute()
    )
    files = results.get("files", [])
    # Skip all Google native formats (Docs, Sheets, Slides, Forms, etc.)
    return [
        f for f in files if not f.get("mimeType", "").startswith(_GOOGLE_APPS_PREFIX)
    ]


def download_file(creds: Credentials, file_id: str) -> bytes:
    import io  # noqa: PLC0415

    from googleapiclient.discovery import build  # noqa: PLC0415
    from googleapiclient.http import MediaIoBaseDownload  # noqa: PLC0415

    service = build("drive", "v3", credentials=creds)
    request = service.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


# ---------- Text extraction (mirrors documents route) ----------

_MAX_TEXT_CHARS = 100_000
_MAX_PDF_PAGES = 50
_MAX_XLSX_ROWS = 200
_MAX_EXTRACT_BYTES = 2 * 1024 * 1024


def _extract_text(content: bytes, filename: str) -> str | None:
    import io  # noqa: PLC0415

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if len(content) > _MAX_EXTRACT_BYTES and ext in ("xlsx", "xls"):
        return None
    text: str | None = None
    try:
        if ext == "pdf":
            from pypdf import PdfReader  # noqa: PLC0415

            reader = PdfReader(io.BytesIO(content))
            pages = reader.pages[:_MAX_PDF_PAGES]
            text = " ".join(page.extract_text() or "" for page in pages)
        elif ext in ("docx", "doc"):
            from docx import Document as DocxDocument  # noqa: PLC0415

            doc = DocxDocument(io.BytesIO(content))
            text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in ("txt", "md", "csv"):
            text = content[:_MAX_TEXT_CHARS].decode("utf-8", errors="ignore")
        elif ext in ("xlsx", "xls"):
            import openpyxl  # noqa: PLC0415

            wb = openpyxl.load_workbook(
                io.BytesIO(content), read_only=True, data_only=True
            )
            parts: list[str] = []
            for ws in wb.worksheets:
                for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    if row_idx >= _MAX_XLSX_ROWS:
                        break
                    parts.extend(str(v) for v in row if v is not None)
            text = " ".join(parts)
    except Exception:
        return None
    if not text:
        return None
    text = text.replace("\x00", "")
    return text[:_MAX_TEXT_CHARS]


# ---------- Tag helper ----------


def _ensure_tag(session: Session, owner_id: uuid.UUID, name: str) -> Tag:
    from sqlmodel import col  # noqa: PLC0415

    tag = session.exec(
        select(Tag).where(Tag.owner_id == owner_id).where(col(Tag.name) == name)
    ).first()
    if not tag:
        tag = Tag(name=name, owner_id=owner_id)
        session.add(tag)
        session.flush()
    return tag


def _link_tag(session: Session, document_id: uuid.UUID, tag_id: uuid.UUID) -> None:
    existing = session.get(DocumentTag, (document_id, tag_id))
    if not existing:
        session.add(DocumentTag(document_id=document_id, tag_id=tag_id))


# ---------- Import orchestration ----------


_DOWNLOAD_WORKERS = 6


def _download_batch(
    creds: Credentials, files: list[dict[str, Any]]
) -> list[tuple[dict[str, Any], bytes | None, str | None]]:
    """Download all files in parallel. Returns (file_meta, content, error) tuples."""
    from concurrent.futures import ThreadPoolExecutor, as_completed  # noqa: PLC0415

    results: list[tuple[dict[str, Any], bytes | None, str | None]] = [
        (f, None, None) for f in files
    ]
    index_map = {f["id"]: i for i, f in enumerate(files)}

    def _dl(f: dict[str, Any]) -> tuple[str, bytes | None, str | None]:
        try:
            return f["id"], download_file(creds, f["id"]), None
        except Exception as exc:
            return f["id"], None, str(exc)

    with ThreadPoolExecutor(max_workers=_DOWNLOAD_WORKERS) as pool:
        futures = {pool.submit(_dl, f): f for f in files}
        for future in as_completed(futures):
            file_id, content, err = future.result()
            idx = index_map[file_id]
            results[idx] = (files[idx], content, err)

    return results


def run_import_job(job_id: uuid.UUID) -> None:
    """Background task — runs after HTTP response is sent."""
    with Session(engine) as session:
        job = session.get(DriveImportJob, job_id)
        if not job:
            return

        connection = session.exec(
            select(DriveConnection).where(DriveConnection.user_id == job.user_id)
        ).first()
        if not connection:
            job.status = "failed"
            job.error_message = "Drive not connected"
            session.add(job)
            session.commit()
            return

        job.status = "running"
        session.add(job)
        session.commit()

        try:
            creds = build_credentials(connection)
            files = list_importable_files(creds, job.folder_id)
            job.total_files = len(files)
            session.add(job)
            session.commit()

            drive_tag = _ensure_tag(session, job.user_id, "from-drive")
            session.commit()

            # Phase 1: download all files in parallel
            downloaded = _download_batch(creds, files)

            # Phase 2: process sequentially (session not thread-safe)
            for file_meta, content, download_err in downloaded:
                filename: str = file_meta["name"]
                try:
                    if download_err or content is None:
                        job.failed_files += 1
                        session.add(job)
                        session.commit()
                        continue
                    sha256 = hashlib.sha256(content).hexdigest()

                    # Duplicate check — skip if this user already has this exact file
                    dup = session.exec(
                        select(DocumentVersion)
                        .join(Document, DocumentVersion.document_id == Document.id)  # type: ignore[arg-type]
                        .where(Document.owner_id == job.user_id)
                        .where(DocumentVersion.sha256 == sha256)
                    ).first()
                    if dup:
                        job.skipped_files += 1
                        session.add(job)
                        session.commit()
                        continue

                    # Derive metadata from filename
                    ext = (
                        filename.rsplit(".", 1)[-1].lower()
                        if "." in filename
                        else "unknown"
                    )
                    raw_title = filename.rsplit(".", 1)[0] if "." in filename else filename
                    title = raw_title or filename  # dotfiles like .env → use full filename

                    # Check for existing document with same title (versioning)
                    existing_doc = session.exec(
                        select(Document)
                        .where(Document.owner_id == job.user_id)
                        .where(Document.title == title)
                        .where(Document.format == ext)
                    ).first()

                    if existing_doc:
                        from sqlmodel import func  # noqa: PLC0415

                        max_v = (
                            session.exec(
                                select(func.max(DocumentVersion.version_number)).where(
                                    DocumentVersion.document_id == existing_doc.id
                                )
                            ).first()
                            or 0
                        )
                        next_v = max_v + 1
                        file_path = (
                            UPLOAD_DIR / f"{existing_doc.id}_v{next_v}_{filename}"
                        )
                        file_path.write_bytes(content)
                        version = DocumentVersion(
                            document_id=existing_doc.id,
                            version_number=next_v,
                            sha256=sha256,
                            file_path=str(file_path),
                            original_filename=filename,
                            file_size=len(content),
                            content_text=_extract_text(content, filename),
                        )
                        session.add(version)
                        doc = existing_doc
                    else:
                        doc = Document(
                            title=title,
                            creator="Google Drive",
                            format=ext,
                            owner_id=job.user_id,
                        )
                        session.add(doc)
                        session.flush()
                        file_path = UPLOAD_DIR / f"{doc.id}_v1_{filename}"
                        file_path.write_bytes(content)
                        version = DocumentVersion(
                            document_id=doc.id,
                            version_number=1,
                            sha256=sha256,
                            file_path=str(file_path),
                            original_filename=filename,
                            file_size=len(content),
                            content_text=_extract_text(content, filename),
                        )
                        session.add(version)

                    session.flush()
                    _link_tag(session, doc.id, drive_tag.id)
                    job.imported_files += 1
                    session.add(job)
                    session.commit()

                except Exception as exc:
                    session.rollback()
                    job.failed_files += 1
                    session.add(job)
                    session.commit()
                    # Re-fetch drive_tag after rollback
                    drive_tag = _ensure_tag(session, job.user_id, "from-drive")
                    session.commit()
                    _ = exc  # individual file failure — continue

            job.status = "completed"
            job.completed_at = datetime.now(timezone.utc)
            session.add(job)
            session.commit()

        except Exception as exc:
            job.status = "failed"
            job.error_message = str(exc)[:2000]
            job.completed_at = datetime.now(timezone.utc)
            session.add(job)
            session.commit()
